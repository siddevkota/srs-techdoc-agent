import io
from pathlib import Path
from typing import Union, BinaryIO

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None


class SRSLoader:
    """Load and extract text from SRS documents in various formats."""
    
    @staticmethod
    def load_from_file(file_path: Union[str, Path]) -> str:
        """
        Load SRS content from a file path.
        
        Args:
            file_path: Path to the SRS file
        
        Returns:
            Extracted text content
        
        Raises:
            ValueError: If file format is unsupported
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return SRSLoader._load_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return SRSLoader._load_docx(file_path)
        elif suffix in ['.txt', '.md']:
            return SRSLoader._load_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    @staticmethod
    def load_from_uploaded_file(uploaded_file) -> str:
        """
        Load SRS content from Streamlit uploaded file object.
        
        Args:
            uploaded_file: Streamlit UploadedFile object
        
        Returns:
            Extracted text content
        """
        file_name = uploaded_file.name.lower()
        file_bytes = uploaded_file.read()
        
        if file_name.endswith('.pdf'):
            return SRSLoader._load_pdf_from_bytes(file_bytes)
        elif file_name.endswith('.docx') or file_name.endswith('.doc'):
            return SRSLoader._load_docx_from_bytes(file_bytes)
        elif file_name.endswith('.txt') or file_name.endswith('.md'):
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_name}")
    
    @staticmethod
    def _load_pdf(file_path: Path) -> str:
        """Extract text from PDF file."""
        if pdfplumber is None:
            raise ImportError("pdfplumber is required for PDF support. Install with: pip install pdfplumber")
        
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _load_pdf_from_bytes(file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        if pdfplumber is None:
            raise ImportError("pdfplumber is required for PDF support. Install with: pip install pdfplumber")
        
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _load_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            return "\n\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX: {str(e)}")
    
    @staticmethod
    def _load_docx_from_bytes(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            return "\n\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX: {str(e)}")
    
    @staticmethod
    def _load_text(file_path: Path) -> str:
        """Load plain text or markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    @staticmethod
    def get_text_stats(text: str) -> dict:
        """Get statistics about extracted text."""
        lines = text.split('\n')
        words = text.split()
        
        return {
            "char_count": len(text),
            "word_count": len(words),
            "line_count": len(lines),
            "non_empty_lines": len([line for line in lines if line.strip()]),
        }
